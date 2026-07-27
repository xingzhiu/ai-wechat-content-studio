import json
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from sqlalchemy import func, select

from app.database import SessionLocal
from app.models import Article, ArticleStatus, ArticleVersion, Citation, FeedItem, Source
from app.website_import import (
    WebsiteImportError,
    _validate_public_url,
    create_article_from_website,
    extract_uploaded_file,
)


def test_private_website_is_rejected():
    with pytest.raises(WebsiteImportError, match="内部服务"):
        _validate_public_url("http://localhost/article")


def test_uploaded_markdown_is_extracted():
    source = extract_uploaded_file(
        "article.md",
        ("# 标题\n\n这是一段用于生成公众号文章的文件正文。" * 20).encode("utf-8"),
    )
    assert source["title"] == "article"
    assert "文件正文" in source["body"]
    assert source["url"] == ""


def test_unsupported_upload_is_rejected():
    with pytest.raises(WebsiteImportError, match="仅支持"):
        extract_uploaded_file("archive.zip", b"not-an-article")


def test_uploaded_json_is_extracted():
    payload = {
        "title": "JSON 示例文章",
        "sections": [
            {"heading": "发生了什么", "content": "这是一段足够长的结构化文章素材。" * 8},
            {"heading": "为什么重要", "content": "字段层级会被保留，方便模型理解上下文。"},
        ],
    }
    source = extract_uploaded_file("article.json", json.dumps(payload, ensure_ascii=False).encode("utf-8"))
    assert source["title"] == "article"
    assert '"sections"' in source["body"]
    assert "JSON 示例文章" in source["body"]


def test_invalid_json_is_rejected():
    with pytest.raises(WebsiteImportError, match="JSON 文件格式无效"):
        extract_uploaded_file("broken.json", b'{"title": }')


def test_docx_with_doc_extension_is_supported():
    document = Document()
    document.add_heading("旧扩展名兼容测试", level=1)
    document.add_paragraph("这是一段用于验证 DOC 扩展名兼容性的文章正文。" * 12)
    buffer = BytesIO()
    document.save(buffer)
    source = extract_uploaded_file("legacy-name.doc", buffer.getvalue())
    assert "旧扩展名兼容测试" in source["body"]


def test_invalid_legacy_doc_is_rejected():
    with pytest.raises(WebsiteImportError, match="旧版 DOC 文件无法解析"):
        extract_uploaded_file("broken.doc", b"not-a-word-document" * 20)


def test_website_creates_pending_article(reset_db, monkeypatch):
    source_data = {
        "url": "https://example.com/article",
        "title": "示例网站文章",
        "author": "Example",
        "body": "这是一段经过提取的网页正文。" * 30,
        "published_at": None,
    }
    generated = {
        "titles": ["主标题", "备选标题一", "备选标题二"],
        "lead": "我从这篇网页中看到一个具体变化。",
        "fact_summary": [
            {"fact": "网页明确写出的事实", "evidence_id": "website-1", "source_url": source_data["url"]}
        ],
        "value_interpretation": "这一变化的价值需要结合实际需求判断。",
        "impact_on_general_users": "普通用户可以先核对自身需求。",
        "developer_focus": ["核对技术边界"],
        "action_recommendations": ["阅读原始网页"],
        "sources": [],
    }

    class FakeResponses:
        def create(self, **_kwargs):
            return SimpleNamespace(output_text=json.dumps(generated, ensure_ascii=False), output=[], status="completed")

    monkeypatch.setattr("app.website_import.fetch_website", lambda _url: source_data)
    monkeypatch.setattr("app.website_import._client", lambda: SimpleNamespace(responses=FakeResponses()))

    with SessionLocal() as db:
        article, reused = create_article_from_website(db, source_data["url"])
        assert reused is False
        assert article.status == ArticleStatus.pending
        assert article.title == "主标题"
        assert article.content["sources"][0]["url"] == source_data["url"]
        assert db.scalar(select(func.count()).select_from(FeedItem)) == 0
        assert db.scalar(select(func.count()).select_from(Citation)) == 0
        assert db.scalar(select(func.count()).select_from(ArticleVersion)) == 1
        assert db.scalar(select(func.count()).select_from(Source).where(Source.kind == "website_import")) == 0

        same_article, reused = create_article_from_website(db, source_data["url"])
        assert reused is True
        assert same_article.id == article.id
        assert db.scalar(select(func.count()).select_from(Article)) == 1

    headers = {"X-Admin-Password": "test"}
    settings_response = reset_db.get("/api/workflows/1/settings", headers=headers).json()
    assert all(item["kind"] != "website_import" for item in settings_response["sources"])
    summary = reset_db.get("/api/database/summary", headers=headers).json()
    assert summary["sources"] == 6
    assert summary["feed_items"] == 0
    assert reset_db.get("/api/database/sources", headers=headers).json()["total"] == 6
    assert reset_db.get("/api/database/feed_items", headers=headers).json()["total"] == 0
